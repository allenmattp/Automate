import docx

doc = docx.Document("demo.docx")

print(len(doc.paragraphs))
print(doc.paragraphs[0].text)
print(doc.paragraphs[1].text)
print(len(doc.paragraphs[1].runs))

for r in doc.paragraphs[1].runs:
    print(r.text)